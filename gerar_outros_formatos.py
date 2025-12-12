"""
Gera documentos de teste em diferentes formatos (Imagem e Word)
"""
import os
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = "contratos_teste"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def criar_imagem_contrato(filename, dados):
    """Cria uma imagem simulando um documento escaneado."""
    # Cria imagem branca
    img = Image.new('RGB', (800, 1100), color='white')
    draw = ImageDraw.Draw(img)
    
    # Usa fonte padrao
    try:
        font_titulo = ImageFont.truetype("arial.ttf", 24)
        font_normal = ImageFont.truetype("arial.ttf", 14)
        font_bold = ImageFont.truetype("arialbd.ttf", 14)
    except:
        font_titulo = ImageFont.load_default()
        font_normal = ImageFont.load_default()
        font_bold = ImageFont.load_default()
    
    y = 50
    
    # Titulo
    draw.text((300, y), dados['titulo'], fill='black', font=font_titulo)
    y += 60
    
    # Linha
    draw.line([(50, y), (750, y)], fill='gray', width=1)
    y += 30
    
    # Conteudo
    for linha in dados['linhas']:
        draw.text((50, y), linha, fill='black', font=font_normal)
        y += 25
    
    # Salva
    filepath = os.path.join(OUTPUT_DIR, filename)
    img.save(filepath, quality=95)
    print(f"Criado: {filename}")


def criar_word_contrato(filename, dados):
    """Cria um documento Word."""
    doc = Document()
    
    # Titulo
    titulo = doc.add_heading(dados['titulo'], 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Paragrafos
    for paragrafo in dados['paragrafos']:
        p = doc.add_paragraph(paragrafo)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Tabela se houver
    if 'tabela' in dados:
        table = doc.add_table(rows=len(dados['tabela']), cols=len(dados['tabela'][0]))
        table.style = 'Table Grid'
        for i, row in enumerate(dados['tabela']):
            for j, cell in enumerate(row):
                table.rows[i].cells[j].text = cell
    
    # Assinatura
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph(f"{dados['cidade']}, {dados['data']}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph("_" * 40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(dados['assinatura'])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    print(f"Criado: {filename}")


if __name__ == "__main__":
    print("=" * 50)
    print("GERANDO DOCUMENTOS EM OUTROS FORMATOS")
    print("=" * 50)
    print()
    
    # ===== IMAGENS (simulando documentos escaneados) =====
    
    criar_imagem_contrato("recibo_escaneado.png", {
        'titulo': 'RECIBO',
        'linhas': [
            '',
            'Recebi de: Empresa ABC Comercio Ltda',
            'CNPJ: 12.345.678/0001-90',
            '',
            'A quantia de: R$ 3.500,00',
            '(Tres mil e quinhentos reais)',
            '',
            'Referente a: Servicos de consultoria',
            'prestados no mes de dezembro/2024',
            '',
            'Data: 10/12/2024',
            '',
            '',
            'Recebedor: Joao Carlos Silva',
            'CPF: 123.456.789-00',
            '',
            '',
            '________________________________',
            'Assinatura'
        ]
    })
    
    criar_imagem_contrato("nota_fiscal_escaneada.jpg", {
        'titulo': 'NOTA FISCAL',
        'linhas': [
            '',
            'EMITENTE: Loja Tech Informatica ME',
            'CNPJ: 98.765.432/0001-11',
            'Endereco: Rua das Flores, 123, Centro',
            '',
            'CLIENTE: Maria Aparecida Santos',
            'CPF: 987.654.321-00',
            '',
            'PRODUTOS:',
            '1x Notebook Dell Inspiron - R$ 4.200,00',
            '1x Mouse Logitech - R$ 150,00',
            '1x Teclado Mecanico - R$ 350,00',
            '',
            'SUBTOTAL: R$ 4.700,00',
            'DESCONTO: R$ 200,00',
            'TOTAL: R$ 4.500,00',
            '',
            'Forma de pagamento: Cartao de Credito',
            'Data: 15/12/2024',
            'NF: 001234'
        ]
    })
    
    # ===== DOCUMENTOS WORD =====
    
    criar_word_contrato("declaracao.docx", {
        'titulo': 'DECLARACAO',
        'paragrafos': [
            '',
            'Eu, Carlos Eduardo Ferreira, brasileiro, solteiro, portador do RG 12.345.678-9 '
            'e CPF 111.222.333-44, residente a Rua das Palmeiras, 456, Jardim Europa, '
            'Sao Paulo/SP, CEP 01234-567, DECLARO para os devidos fins que:',
            '',
            'Trabalho na empresa Tech Solutions S.A., CNPJ 55.666.777/0001-88, '
            'desde 01/03/2020, exercendo a funcao de Analista de Sistemas, '
            'com salario mensal de R$ 8.500,00 (oito mil e quinhentos reais).',
            '',
            'Declaro ainda que as informacoes acima sao verdadeiras e assumo '
            'inteira responsabilidade pelas mesmas.',
        ],
        'cidade': 'Sao Paulo/SP',
        'data': '12/12/2024',
        'assinatura': 'Carlos Eduardo Ferreira'
    })
    
    criar_word_contrato("orcamento.docx", {
        'titulo': 'ORCAMENTO DE SERVICOS',
        'paragrafos': [
            '',
            'Cliente: Empresa Delta Comercio Ltda',
            'CNPJ: 33.444.555/0001-66',
            'Contato: Sr. Roberto Almeida',
            'Email: roberto@deltaltda.com.br',
            'Telefone: (11) 98765-4321',
            '',
            'Prezado cliente, segue orcamento conforme solicitado:',
            '',
        ],
        'tabela': [
            ['Item', 'Descricao', 'Qtd', 'Valor Unit.', 'Total'],
            ['1', 'Desenvolvimento de Website', '1', 'R$ 5.000,00', 'R$ 5.000,00'],
            ['2', 'Hospedagem Anual', '1', 'R$ 600,00', 'R$ 600,00'],
            ['3', 'Manutencao Mensal', '12', 'R$ 300,00', 'R$ 3.600,00'],
            ['', '', '', 'TOTAL:', 'R$ 9.200,00'],
        ],
        'cidade': 'Sao Paulo/SP',
        'data': '12/12/2024',
        'assinatura': 'Agencia Web Solutions'
    })
    
    criar_word_contrato("termo_responsabilidade.docx", {
        'titulo': 'TERMO DE RESPONSABILIDADE',
        'paragrafos': [
            '',
            'Eu, Ana Paula Oliveira, CPF 222.333.444-55, funcionaria da empresa '
            'ABC Tecnologia Ltda, matricula 12345, declaro ter recebido os seguintes '
            'equipamentos para uso profissional:',
            '',
            '- 1 (um) Notebook Dell Latitude, Patrimonio: NB-2024-0123',
            '- 1 (um) Celular iPhone 14, Patrimonio: CEL-2024-0456',
            '- 1 (um) Headset Jabra, Patrimonio: HS-2024-0789',
            '',
            'Comprometo-me a zelar pelos equipamentos acima descritos, utilizando-os '
            'exclusivamente para fins profissionais, e a devolve-los em perfeito estado '
            'quando solicitado ou em caso de desligamento da empresa.',
            '',
            'Valor total dos equipamentos: R$ 12.500,00',
        ],
        'cidade': 'Sao Paulo/SP',
        'data': '01/12/2024',
        'assinatura': 'Ana Paula Oliveira'
    })
    
    print()
    print("=" * 50)
    print("Arquivos gerados:")
    for f in os.listdir(OUTPUT_DIR):
        ext = f.split('.')[-1].upper()
        print(f"  {ext}: {f}")
    print("=" * 50)

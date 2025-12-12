"""
=============================================================================
CONVERSOR - Suporte a multiplos formatos de documento
=============================================================================

Este modulo converte diferentes formatos de documento para PDF ou extrai
texto diretamente, permitindo processamento unificado.

Formatos suportados:
- PDF: Processamento direto
- Imagens (PNG, JPG, JPEG, BMP, TIFF): Conversao para PDF ou OCR direto
- Word (DOCX, DOC): Conversao para PDF ou extracao de texto
"""

import os
import sys
from io import BytesIO
from typing import Tuple, List, Optional
from pathlib import Path

# Tenta importar dependencias
try:
    from PIL import Image
    PIL_DISPONIVEL = True
except ImportError:
    PIL_DISPONIVEL = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_DISPONIVEL = True
except ImportError:
    PYMUPDF_DISPONIVEL = False

try:
    from docx import Document
    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False

# =============================================================================
# CONSTANTES
# =============================================================================

FORMATOS_PDF = ['.pdf']
FORMATOS_IMAGEM = ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif']
FORMATOS_WORD = ['.docx', '.doc']
TODOS_FORMATOS = FORMATOS_PDF + FORMATOS_IMAGEM + FORMATOS_WORD


def get_extensao(filename: str) -> str:
    """Retorna a extensao do arquivo em minusculo."""
    return Path(filename).suffix.lower()


def eh_pdf(filename: str) -> bool:
    """Verifica se o arquivo e PDF."""
    return get_extensao(filename) in FORMATOS_PDF


def eh_imagem(filename: str) -> bool:
    """Verifica se o arquivo e uma imagem."""
    return get_extensao(filename) in FORMATOS_IMAGEM


def eh_word(filename: str) -> bool:
    """Verifica se o arquivo e Word."""
    return get_extensao(filename) in FORMATOS_WORD


def formato_suportado(filename: str) -> bool:
    """Verifica se o formato e suportado."""
    return get_extensao(filename) in TODOS_FORMATOS


# =============================================================================
# CONVERSAO DE IMAGEM PARA PDF
# =============================================================================

def imagem_para_pdf(image_file, filename: str = "image.png") -> BytesIO:
    """
    Converte uma imagem para PDF.

    Args:
        image_file: Arquivo de imagem (bytes ou file-like object)
        filename: Nome do arquivo original

    Returns:
        BytesIO contendo o PDF gerado
    """
    if not PIL_DISPONIVEL:
        raise ImportError("Pillow nao esta instalado. Execute: pip install Pillow")

    if not PYMUPDF_DISPONIVEL:
        raise ImportError("PyMuPDF nao esta instalado. Execute: pip install pymupdf")

    # Le a imagem
    if hasattr(image_file, 'read'):
        image_file.seek(0)
        img_bytes = image_file.read()
        image_file.seek(0)
    else:
        img_bytes = image_file

    # Abre com PIL para garantir formato correto
    img = Image.open(BytesIO(img_bytes))

    # Converte para RGB se necessario (ex: PNG com transparencia)
    if img.mode in ('RGBA', 'P'):
        img = img.convert('RGB')

    # Salva como PDF usando PIL
    pdf_buffer = BytesIO()
    img.save(pdf_buffer, 'PDF', resolution=100.0)
    pdf_buffer.seek(0)

    return pdf_buffer


def extrair_texto_imagem(image_file, idioma: str = "por+eng") -> Tuple[str, List[dict], Tuple[float, float]]:
    """
    Extrai texto de uma imagem usando OCR.

    Args:
        image_file: Arquivo de imagem
        idioma: Idioma(s) para OCR

    Returns:
        - texto_completo: Todo o texto extraido
        - palavras: Lista de dicts com coordenadas
        - image_size: Tupla (largura, altura)
    """
    try:
        import pytesseract
        from ocr_engine import verificar_tesseract_instalado, configurar_tesseract_windows
    except ImportError:
        raise ImportError("pytesseract nao esta instalado. Execute: pip install pytesseract")

    if not verificar_tesseract_instalado():
        if sys.platform == 'win32':
            configurar_tesseract_windows()
        if not verificar_tesseract_instalado():
            raise RuntimeError("Tesseract nao esta instalado no sistema")

    # Le a imagem
    if hasattr(image_file, 'read'):
        image_file.seek(0)
        img_bytes = image_file.read()
        image_file.seek(0)
    else:
        img_bytes = image_file

    img = Image.open(BytesIO(img_bytes))
    img_width, img_height = img.size

    # Executa OCR
    dados_ocr = pytesseract.image_to_data(
        img,
        lang=idioma,
        output_type=pytesseract.Output.DICT,
        config='--psm 6'
    )

    palavras = []
    texto_linhas = []

    n_boxes = len(dados_ocr['text'])

    for i in range(n_boxes):
        texto = dados_ocr['text'][i].strip()
        conf = int(dados_ocr['conf'][i])

        if not texto or conf < 30:
            continue

        x = dados_ocr['left'][i]
        y = dados_ocr['top'][i]
        w = dados_ocr['width'][i]
        h = dados_ocr['height'][i]

        palavras.append({
            "text": texto,
            "x0": x,
            "top": y,
            "x1": x + w,
            "bottom": y + h,
            "width": w,
            "height": h,
            "confidence": conf
        })

        texto_linhas.append(texto)

    texto_completo = " ".join(texto_linhas)

    return texto_completo, palavras, (img_width, img_height)


# =============================================================================
# CONVERSAO DE WORD PARA PDF/TEXTO
# =============================================================================

def extrair_texto_docx(docx_file) -> Tuple[str, List[dict], Tuple[float, float]]:
    """
    Extrai texto de um arquivo DOCX.

    Args:
        docx_file: Arquivo DOCX

    Returns:
        - texto_completo: Todo o texto extraido
        - palavras: Lista de dicts (sem coordenadas precisas para DOCX)
        - page_size: Tupla (largura, altura) estimada
    """
    if not DOCX_DISPONIVEL:
        raise ImportError("python-docx nao esta instalado. Execute: pip install python-docx")

    # Le o documento
    if hasattr(docx_file, 'read'):
        docx_file.seek(0)
        doc = Document(docx_file)
        docx_file.seek(0)
    else:
        doc = Document(BytesIO(docx_file))

    # Extrai texto de todos os paragrafos
    texto_completo = ""
    palavras = []

    # Posicao Y simulada (DOCX nao tem coordenadas reais)
    y_pos = 50
    x_pos = 50

    for para in doc.paragraphs:
        texto_para = para.text.strip()
        if texto_para:
            texto_completo += texto_para + "\n"

            # Divide em palavras e cria coordenadas simuladas
            for palavra in texto_para.split():
                largura_palavra = len(palavra) * 7  # Estimativa

                palavras.append({
                    "text": palavra,
                    "x0": x_pos,
                    "top": y_pos,
                    "x1": x_pos + largura_palavra,
                    "bottom": y_pos + 12,
                    "width": largura_palavra,
                    "height": 12
                })

                x_pos += largura_palavra + 5

                # Quebra de linha simulada
                if x_pos > 500:
                    x_pos = 50
                    y_pos += 15

            y_pos += 20
            x_pos = 50

    # Extrai texto das tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto_cell = cell.text.strip()
                if texto_cell:
                    texto_completo += texto_cell + " "

                    for palavra in texto_cell.split():
                        largura_palavra = len(palavra) * 7

                        palavras.append({
                            "text": palavra,
                            "x0": x_pos,
                            "top": y_pos,
                            "x1": x_pos + largura_palavra,
                            "bottom": y_pos + 12,
                            "width": largura_palavra,
                            "height": 12
                        })

                        x_pos += largura_palavra + 5

            y_pos += 15
            x_pos = 50

    # Tamanho de pagina A4 simulado
    page_size = (595, 842)

    return texto_completo.strip(), palavras, page_size


def docx_para_pdf(docx_file) -> BytesIO:
    """
    Converte DOCX para PDF.

    Nota: Esta conversao e simplificada. Para conversao perfeita,
    seria necessario usar LibreOffice ou Microsoft Word.

    Args:
        docx_file: Arquivo DOCX

    Returns:
        BytesIO contendo o PDF gerado
    """
    if not PYMUPDF_DISPONIVEL:
        raise ImportError("PyMuPDF nao esta instalado. Execute: pip install pymupdf")

    if not DOCX_DISPONIVEL:
        raise ImportError("python-docx nao esta instalado. Execute: pip install python-docx")

    # Extrai texto do DOCX
    texto, _, _ = extrair_texto_docx(docx_file)

    # Cria PDF com o texto
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4

    # Insere texto no PDF
    text_rect = fitz.Rect(50, 50, 545, 792)
    page.insert_textbox(
        text_rect,
        texto,
        fontsize=11,
        fontname="helv"
    )

    # Salva em buffer
    pdf_buffer = BytesIO()
    doc.save(pdf_buffer)
    doc.close()

    pdf_buffer.seek(0)
    return pdf_buffer


# =============================================================================
# FUNCAO PRINCIPAL DE CONVERSAO
# =============================================================================

def processar_documento(file, filename: str) -> Tuple[BytesIO, str, bool]:
    """
    Processa qualquer documento suportado e retorna em formato utilizavel.

    Args:
        file: Arquivo (bytes ou file-like object)
        filename: Nome do arquivo com extensao

    Returns:
        - pdf_file: Arquivo em formato PDF (BytesIO)
        - tipo_original: Tipo do arquivo original
        - convertido: True se foi convertido, False se ja era PDF
    """
    extensao = get_extensao(filename)

    if hasattr(file, 'read'):
        file.seek(0)
        conteudo = file.read()
        file.seek(0)
    else:
        conteudo = file

    if eh_pdf(filename):
        # Ja e PDF, retorna direto
        return BytesIO(conteudo), "PDF", False

    elif eh_imagem(filename):
        # Converte imagem para PDF
        pdf_buffer = imagem_para_pdf(conteudo, filename)
        return pdf_buffer, "Imagem", True

    elif eh_word(filename):
        # Converte Word para PDF
        pdf_buffer = docx_para_pdf(BytesIO(conteudo))
        return pdf_buffer, "Word", True

    else:
        raise ValueError(f"Formato nao suportado: {extensao}")


def extrair_texto_documento(
    file,
    filename: str,
    forcar_ocr: bool = False
) -> Tuple[str, List[dict], Tuple[float, float], str]:
    """
    Extrai texto de qualquer documento suportado.

    Args:
        file: Arquivo
        filename: Nome do arquivo
        forcar_ocr: Forcar uso de OCR

    Returns:
        - texto: Texto extraido
        - palavras: Lista com coordenadas
        - page_size: Tamanho da pagina
        - metodo: Metodo usado
    """
    extensao = get_extensao(filename)

    if eh_pdf(filename):
        # Usa o OCR engine existente
        from ocr_engine import extrair_texto_automatico
        return extrair_texto_automatico(file, forcar_ocr=forcar_ocr)

    elif eh_imagem(filename):
        # Extrai direto da imagem com OCR
        texto, palavras, size = extrair_texto_imagem(file)
        return texto, palavras, size, "OCR (imagem)"

    elif eh_word(filename):
        # Extrai do DOCX
        texto, palavras, size = extrair_texto_docx(file)
        return texto, palavras, size, "Word (texto)"

    else:
        raise ValueError(f"Formato nao suportado: {extensao}")


# =============================================================================
# TESTE DO MODULO
# =============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("TESTE DO CONVERSOR DE DOCUMENTOS")
    print("=" * 60)

    print("\n[1] Verificando dependencias...")
    print(f"    - Pillow: {'OK' if PIL_DISPONIVEL else 'NAO INSTALADO'}")
    print(f"    - PyMuPDF: {'OK' if PYMUPDF_DISPONIVEL else 'NAO INSTALADO'}")
    print(f"    - python-docx: {'OK' if DOCX_DISPONIVEL else 'NAO INSTALADO'}")

    print("\n[2] Formatos suportados:")
    print(f"    - PDF: {FORMATOS_PDF}")
    print(f"    - Imagens: {FORMATOS_IMAGEM}")
    print(f"    - Word: {FORMATOS_WORD}")

    print("\n[3] Testando deteccao de formato...")
    testes = ["documento.pdf", "foto.jpg", "IMAGEM.PNG", "contrato.docx", "arquivo.doc"]
    for teste in testes:
        tipo = "PDF" if eh_pdf(teste) else "Imagem" if eh_imagem(teste) else "Word" if eh_word(teste) else "?"
        print(f"    - {teste}: {tipo}")

    print("\n" + "=" * 60)
    print("CONVERSOR FUNCIONANDO!")
    print("=" * 60)

"""
=============================================================================
CONFTEST - Fixtures compartilhadas para os testes
=============================================================================
"""

import pytest
import sys
import os
from io import BytesIO
from pathlib import Path

# Adiciona o diretorio raiz ao path para importar os modulos
sys.path.insert(0, str(Path(__file__).parent.parent))

from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4


# =============================================================================
# FIXTURES DE ARQUIVOS DE TESTE
# =============================================================================

@pytest.fixture
def pdf_simples() -> BytesIO:
    """
    Gera um PDF simples com texto para testes.
    Contém campos típicos como nome, CPF, data e valor.
    """
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    
    # Adiciona conteúdo
    c.setFont("Helvetica-Bold", 16)
    c.drawString(200, 750, "DOCUMENTO DE TESTE")
    
    c.setFont("Helvetica", 12)
    c.drawString(50, 700, "Nome: João da Silva")
    c.drawString(50, 680, "CPF: 123.456.789-00")
    c.drawString(50, 660, "Data: 01/01/2024")
    c.drawString(50, 640, "Valor: R$ 1.500,00")
    c.drawString(50, 620, "Email: joao@teste.com")
    c.drawString(50, 600, "Telefone: (11) 99999-9999")
    
    c.save()
    buffer.seek(0)
    return buffer


@pytest.fixture
def pdf_vazio() -> BytesIO:
    """Gera um PDF vazio (sem texto) para testar OCR."""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.save()
    buffer.seek(0)
    return buffer


@pytest.fixture
def imagem_com_texto() -> BytesIO:
    """
    Gera uma imagem PNG com texto para testar OCR.
    """
    # Cria imagem branca
    img = Image.new('RGB', (800, 600), color='white')
    
    # Adiciona texto simples usando PIL
    from PIL import ImageDraw, ImageFont
    draw = ImageDraw.Draw(img)
    
    try:
        # Tenta usar fonte do sistema
        font = ImageFont.truetype("arial.ttf", 24)
    except:
        # Fallback para fonte padrão
        font = ImageFont.load_default()
    
    draw.text((50, 50), "DOCUMENTO DE TESTE", fill='black', font=font)
    draw.text((50, 100), "Nome: Maria Santos", fill='black', font=font)
    draw.text((50, 150), "CPF: 987.654.321-00", fill='black', font=font)
    
    buffer = BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)
    return buffer


@pytest.fixture
def imagem_jpg() -> BytesIO:
    """Gera uma imagem JPG simples."""
    img = Image.new('RGB', (400, 300), color='lightblue')
    buffer = BytesIO()
    img.save(buffer, format='JPEG')
    buffer.seek(0)
    return buffer


@pytest.fixture
def imagem_bmp() -> BytesIO:
    """Gera uma imagem BMP simples."""
    img = Image.new('RGB', (400, 300), color='lightgreen')
    buffer = BytesIO()
    img.save(buffer, format='BMP')
    buffer.seek(0)
    return buffer


@pytest.fixture
def docx_simples(tmp_path) -> BytesIO:
    """
    Gera um documento DOCX simples para testes.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        pytest.skip("python-docx não instalado")
    
    doc = Document()
    doc.add_heading("Contrato de Teste", 0)
    doc.add_paragraph("Nome do Cliente: Pedro Oliveira")
    doc.add_paragraph("CPF: 111.222.333-44")
    doc.add_paragraph("Data: 15/06/2024")
    doc.add_paragraph("Valor Total: R$ 5.000,00")
    
    # Adiciona tabela
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Preço"
    table.cell(1, 0).text = "Serviço A"
    table.cell(1, 1).text = "R$ 2.500,00"
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@pytest.fixture
def temp_dir(tmp_path):
    """Diretório temporário para arquivos de teste."""
    return tmp_path


# =============================================================================
# FIXTURES PARA MOCKS
# =============================================================================

@pytest.fixture
def mock_mapeamentos():
    """Mapeamentos de exemplo para testes de substituição."""
    return [
        {
            "tipo": "NOME_CLIENTE",
            "descricao": "Nome do Cliente",
            "texto_original": "João da Silva",
            "x0": 100,
            "top": 700,
            "x1": 200,
            "bottom": 712
        },
        {
            "tipo": "CPF_CLIENTE",
            "descricao": "CPF do Cliente",
            "texto_original": "123.456.789-00",
            "x0": 100,
            "top": 680,
            "x1": 220,
            "bottom": 692
        },
        {
            "tipo": "DATA_DOCUMENTO",
            "descricao": "Data do Documento",
            "texto_original": "01/01/2024",
            "x0": 100,
            "top": 660,
            "x1": 180,
            "bottom": 672
        },
        {
            "tipo": "VALOR_TOTAL",
            "descricao": "Valor Total",
            "texto_original": "R$ 1.500,00",
            "x0": 100,
            "top": 640,
            "x1": 190,
            "bottom": 652
        }
    ]


@pytest.fixture
def mock_variaveis_llm():
    """Variáveis retornadas pela LLM para testes."""
    return [
        {"valor_original": "João da Silva", "tipo": "NOME_CLIENTE", "descricao": "Nome do Cliente"},
        {"valor_original": "123.456.789-00", "tipo": "CPF_CLIENTE", "descricao": "CPF do Cliente"},
        {"valor_original": "01/01/2024", "tipo": "DATA_DOCUMENTO", "descricao": "Data do Documento"},
        {"valor_original": "R$ 1.500,00", "tipo": "VALOR_TOTAL", "descricao": "Valor Total"}
    ]


@pytest.fixture
def mock_palavras():
    """Lista de palavras com coordenadas para testes de mapeamento."""
    return [
        {"text": "João", "x0": 100, "top": 700, "x1": 130, "bottom": 712},
        {"text": "da", "x0": 135, "top": 700, "x1": 150, "bottom": 712},
        {"text": "Silva", "x0": 155, "top": 700, "x1": 200, "bottom": 712},
        {"text": "123.456.789-00", "x0": 100, "top": 680, "x1": 220, "bottom": 692},
        {"text": "01/01/2024", "x0": 100, "top": 660, "x1": 180, "bottom": 672},
        {"text": "R$", "x0": 100, "top": 640, "x1": 120, "bottom": 652},
        {"text": "1.500,00", "x0": 125, "top": 640, "x1": 190, "bottom": 652}
    ]
